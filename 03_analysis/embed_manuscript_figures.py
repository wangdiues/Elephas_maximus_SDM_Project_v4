from __future__ import annotations

import os
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches


ROOT = Path(__file__).parent.parent / "05_manuscripts" / "Ecology and Evolution"
MAIN_DOCS = [
    ROOT / "Ecology_and_Evolution_REVISED.docx",
    ROOT / "Submission" / "Main" / "Ecology_and_Evolution_REVISED.docx",
]
SUPP_DOCS = [
    ROOT / "Ecology_and_Evolution_SUPPLEMENTARY_REVISED.docx",
    ROOT / "Submission" / "Supplementary" / "Ecology_and_Evolution_SUPPLEMENTARY_REVISED.docx",
]

MAIN_FIGS = [
    ROOT / "figures" / "main" / "Figure_01_Study_area_and_spatial_cross_validation_design.png",
    ROOT / "figures" / "main" / "Figure_02_Environmental_drivers_of_Elephas_maximus_habitat_suitability_in_Bhutan.png",
    ROOT / "figures" / "main" / "Figure_03_Present_day_habitat_suitability_of_Elephas_maximus_in_Bhutan.png",
    ROOT / "figures" / "main" / "Figure_04_Future_habitat_trajectories_for_Elephas_maximus_in_Bhutan_across_four_SSPs_2021_2100.png",
    ROOT / "figures" / "main" / "Figure_05_Climate_refugia_stability_classification_for_Elephas_maximus_in_Bhutan.png",
    ROOT / "figures" / "main" / "Figure_06_Present_day_human_elephant_conflict_risk_in_Bhutan.png",
]

SUPP_FIGS = [
    ROOT / "figures" / "supplementary" / "Figure_S01_Algorithm_specific_present_day_habitat_suitability_maps_for_Elephas_maximus_in_Bhutan.png",
    ROOT / "figures" / "supplementary" / "Figure_S02_Receiver_operating_characteristic_ROC_curves_for_all_four_algorithms.png",
    ROOT / "figures" / "supplementary" / "Figure_S03_Calibration_plots_for_all_four_algorithms.png",
    ROOT / "figures" / "supplementary" / "Figure_S04_Null_model_comparison_observed_AUC_versus_permuted_null_distribution.png",
    ROOT / "figures" / "supplementary" / "Figure_S05_Spatial_cross_validation_fold_performance.png",
    ROOT / "figures" / "supplementary" / "Figure_S06_Predictor_collinearity_matrix_all_35_candidate_predictors.png",
    ROOT / "figures" / "supplementary" / "Figure_S07_Threshold_sensitivity_analysis.png",
    ROOT / "figures" / "supplementary" / "Figure_S08_Climate_envelope_analysis.png",
    ROOT / "figures" / "supplementary" / "Figure_S09_Full_marginal_response_curves_all_17_predictors_4_algorithms.png",
    ROOT / "figures" / "supplementary" / "Figure_S10a_Occurrence_and_absence_data_spatial_distribution.png",
    ROOT / "figures" / "supplementary" / "Figure_S10b_Kernel_density_of_presence_records.png",
    ROOT / "figures" / "supplementary" / "Figure_S11_Inter_GCM_agreement_maps.png",
    ROOT / "figures" / "supplementary" / "Figure_S12_Inter_GCM_spread_maps.png",
    ROOT / "figures" / "supplementary" / "Figure_S13_GCM_reliability_ranking.png",
    ROOT / "figures" / "supplementary" / "Figure_S14_Individual_GCM_habitat_area_trajectories.png",
    ROOT / "figures" / "supplementary" / "Figure_S15a_Future_ensemble_suitability_2021_2050.png",
    ROOT / "figures" / "supplementary" / "Figure_S15b_Future_ensemble_suitability_2051_2080.png",
    ROOT / "figures" / "supplementary" / "Figure_S15c_Future_ensemble_suitability_2071_2100.png",
    ROOT / "figures" / "supplementary" / "Figure_S16a_Habitat_change_map_2021_2050.png",
    ROOT / "figures" / "supplementary" / "Figure_S16b_Habitat_change_map_2051_2080.png",
    ROOT / "figures" / "supplementary" / "Figure_S16c_Habitat_change_map_2071_2100.png",
    ROOT / "figures" / "supplementary" / "Figure_S17a_Gain_loss_persistence_2021_2050.png",
    ROOT / "figures" / "supplementary" / "Figure_S17b_Gain_loss_persistence_2051_2080.png",
    ROOT / "figures" / "supplementary" / "Figure_S17c_Gain_loss_persistence_2071_2100.png",
    ROOT / "figures" / "supplementary" / "Figure_S18a_Ensemble_uncertainty_2021_2050.png",
    ROOT / "figures" / "supplementary" / "Figure_S18b_Ensemble_uncertainty_2051_2080.png",
    ROOT / "figures" / "supplementary" / "Figure_S18c_Ensemble_uncertainty_2071_2100.png",
    ROOT / "figures" / "supplementary" / "Figure_S19a_Future_conflict_risk_2021_2050.png",
    ROOT / "figures" / "supplementary" / "Figure_S19b_Future_conflict_risk_2051_2080.png",
    ROOT / "figures" / "supplementary" / "Figure_S19c_Future_conflict_risk_2071_2100.png",
    ROOT / "figures" / "supplementary" / "Figure_S19d_Conflict_risk_change_2021_2050.png",
    ROOT / "figures" / "supplementary" / "Figure_S19e_Conflict_risk_change_2051_2080.png",
    ROOT / "figures" / "supplementary" / "Figure_S19f_Conflict_risk_change_2071_2100.png",
    ROOT / "figures" / "supplementary" / "Figure_S20_Spatial_residual_autocorrelogram_Morans_I.png",
    ROOT / "figures" / "supplementary" / "Figure_S21_MESS_extrapolation_risk_map.png",
]
def add_picture_after(paragraph, image_path: Path, width_inches: float = 6.6) -> None:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def remove_picture_paragraphs(doc: Document) -> None:
    for para in list(doc.paragraphs):
        xml = para._p.xml
        if "pic:pic" in xml or "w:drawing" in xml:
            if not para.text.strip():
                p = para._p
                p.getparent().remove(p)


def replace_main_doc(doc_path: Path, figure_paths: list[Path]) -> None:
    doc = Document(str(doc_path))
    remove_picture_paragraphs(doc)

    caption_map = {
        1: figure_paths[0],
        2: figure_paths[1],
        3: figure_paths[2],
        4: figure_paths[3],
        5: figure_paths[4],
        6: figure_paths[5],
    }

    inserted = 0
    for para in list(doc.paragraphs):
        txt = para.text.strip()
        for fig_num, img in caption_map.items():
            if txt.startswith(f"Figure {fig_num}."):
                add_picture_after(para, img)
                inserted += 1
                break

    if inserted != len(figure_paths):
        raise RuntimeError(f"Inserted {inserted} main figures, expected {len(figure_paths)}")

    tmp = doc_path.with_suffix(".tmp.docx")
    doc.save(str(tmp))
    os.replace(tmp, doc_path)


def replace_supplementary_media(doc_path: Path, figure_paths: list[Path]) -> None:
    # The document keeps its layout in package XML; we only swap the media payloads.
    with zipfile.ZipFile(doc_path, "r") as zin:
        rels = zin.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")
        targets = re.findall(r'Target="media/(image\d+\.png)"', rels)
        if not targets:
            raise RuntimeError(f"No image relationships found in {doc_path.name}")

        # Two composite figures are reused so that every relationship target receives a replacement.
        sources = list(figure_paths)
        if len(sources) < len(targets):
            sources.extend(
                [
                    ROOT / "figures" / "supplementary" / "Figure_S08_Climate_envelope_analysis.png",
                    ROOT / "figures" / "supplementary" / "Figure_S14_Individual_GCM_habitat_area_trajectories.png",
                ]
            )
        if len(sources) != len(targets):
            raise RuntimeError(
                f"Source/target mismatch for {doc_path.name}: {len(sources)} sources vs {len(targets)} targets"
            )

        tmp = doc_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/media/"):
                    name = Path(item.filename).name
                    try:
                        idx = targets.index(name)
                    except ValueError:
                        pass
                    else:
                        data = sources[idx].read_bytes()
                zout.writestr(item, data)

    os.replace(tmp, doc_path)


def main() -> None:
    for doc_path in MAIN_DOCS:
        if doc_path.exists():
            replace_main_doc(doc_path, MAIN_FIGS)

    for doc_path in SUPP_DOCS:
        if doc_path.exists():
            replace_supplementary_media(doc_path, SUPP_FIGS)

    print("Embedded figures into manuscript and supplementary DOCX files.")


if __name__ == "__main__":
    main()
